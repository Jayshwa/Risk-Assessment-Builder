# region Imports
from datetime import date, datetime
from pydoc import doc
import time
import os
import json
from tkinter import *
from tkinter import font as tkFont
from tkinter import messagebox
from tkinter.scrolledtext import ScrolledText
from tkinter import ttk
import itertools
import textwrap
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

# endregion

# region Global Variables

# Main frame for application
master_window = Tk()

# Colour pallette
white = "#ffffff"
light_grey = "#4a484d"
med_grey = "#2a292b"
dark_grey = "#232323"

# Variables for adding a new row to the risk assessment section
table_row = itertools.count()
table_contents = []

# Variable for wrapping text and widgets that go beyond frame borders
wrap_widget = 0

# Variable for retrieving the current date/time properties relating to 'now'
now = datetime.now()

# Variables defining the file path of:
# - This file
# - Directory where risk assessment word files are stored
# - Directory where risk modules that are created in
#   the new_risk_area section by the user are stored
file_location = os.path.dirname(__file__)
risk_assessments_folder = os.path.join(file_location, "Risk Assessments")
file_location_risks_folder = os.path.join(file_location, "Risks")
templates_folder = os.path.join(file_location, "Template")

# Configuration for adjusting the design of treeview
s = ttk.Style()
s.configure("Treeview", rowheight=140)
s.configure("Treeview.Heading", font=(None, 13))
s.configure("Treeview", font=(None, 12))
s.configure("Treeview", borderwidth=0)
# endregion

# region Generate directories

# Main body of code defined in function that is executed
# -> if __name__ == "__main__"
def main():
    """Changes the current directory to the location of this file
    Allows the creation of files and directories that this
    program depends on relative to it's current location."""
    os.chdir(file_location)

    if not os.path.exists("Risks"):
        """Creates a sub-directory called "Risks" within the directory of this __file__
        Risks that are created by the user will be stored here as txt files
        """
        print(f"Creating directory 'Risks' in {file_location}")
        os.makedirs("Risks")

    if not os.path.exists("Template"):
        """Creates a sub-directory called "Template" within the directory of this __file__
        Contains the default text for the risk assessments.
        """
        print(f"Creating directory 'Template' in {file_location}")
        os.makedirs("Template")

    # endregion

    # region Functions
    # ----------------------------FUNCTIONS-----------------------------
    def side_bar_control():
        """Controls the side bar visibiity"""
        print("Triggered side_bar_control()")
        if side_bar.container_frame.winfo_ismapped():
            # Hides the side_bar if packed
            print("side_bar is visible, hiding side_bar")
            side_bar.container_frame.pack_forget()

        elif not side_bar.container_frame.winfo_ismapped():
            # Displays the side_bar if hidden
            print("side_bar is not visible, packing side_bar")
            main_area.pack_forget()
            side_bar.container_frame.pack(fill=BOTH, expand=False, side=LEFT)
            main_area.pack(fill=BOTH, expand=True, side=LEFT)

    def refresh_sidebar():
        """Destroys all current buttons within the side bar
        Generates new buttons to replace them
        """
        for widget in side_bar.container_area.winfo_children():
            widget.destroy()

        # Variables
        row = 0

        """Creates a current list of files present in the
        file_location_risks_folder directory
        """
        risk_modules = [x for x in os.listdir(file_location_risks_folder)]
        risk_module_names = [x for x in os.listdir(file_location_risks_folder)]

        """List for storing lambda functions created in
        for loop that will then be assigned to
        dynamically created side bar buttons
        """
        lambdas = []

        for i in risk_modules:
            """Function for dynamically creating buttons in the side bar
            after listing the files listed in risk_module_names
            """
            risk_module_names_pop = risk_module_names.pop(0)
            lambdas.append(lambda i=i: display_risk(i))
            lambdas_pop = lambdas.pop(0)
            while row <= len(risk_modules):
                row += 1
                new_row = [].append(row)
                new_button = Risks(
                    side_bar.container_area,
                    risk_module_names_pop.rstrip().replace(".json", ""),
                    row,
                    0,
                    lambdas_pop,
                )
                break

    def open_risk_assessment(pathtofile):
        """Opens the selected word document created by the user
        that is stored within the risk_assessments_folder directory
        """
        print("Triggered open_risk_assessment()")
        # Changes the directory to be searched
        os.chdir(risk_assessments_folder)

        print("Changing directory to: risk_assessments_folder")
        # Opens the file if present
        os.startfile(pathtofile)

        print(f"Opening file: {pathtofile}")
        # Returns to the main directory where __file__ is stored
        os.chdir(file_location)

        print("Changing directory to: file_location")

    def edit_table(e):
        """Edit the row values selected by clicking the rows
        in the risk_assessment_info_container_table widget
        """
        print("Triggered edit_table()")
        row_id = risk_assessment_info_container_table.selection()
        table_data = risk_assessment_info_container_table.item(row_id)["values"]

        """Clear the entry boxes once new information
        has been updated
        """
        update_risk.delete(0, END)
        update_likelihood.delete(0, END)
        update_severity.delete(0, END)
        update_risk.insert(0, table_data[0])
        update_likelihood.insert(0, table_data[1])
        update_severity.insert(0, table_data[2])

    def clear_risks():
        """Deletes all currently added risks from
        the risk_assessment_info_container_table
        """
        print("Triggered clear_risks()")
        if not risk_assessment_area.container_frame.winfo_ismapped():
            return
        else:
            clear_query = messagebox.askyesno(
                "Clear risk assessment",
                "Are you sure you want to clear the current risk assessment?",
            )
            if clear_query == True:
                print("Clearing current risk assessment")
                for child in risk_assessment_info_container_table.get_children():
                    risk_assessment_info_container_table.delete(child)
                table_contents.clear()
                risk_assessment_name.config(state=NORMAL)
            else:
                print("Cancelling clearing current risk assessment")
                pass

    def new_risk_area_control():
        """Opens and closes the new_risk_area
        based on whether or not it is already packed
        """
        print("Triggered new_risk_area_control()")
        current_assessments_area.container_frame.pack_forget()
        risk_assessment_area.container_frame.pack_forget()
        main_title.pack_forget()
        intro_text.pack_forget()

        if new_risk_area.container_frame.winfo_ismapped():
            # Hides the new_risk_area if packed
            print("new_risk_area is visible, hiding new_risk_area")
            side_bar.container_frame.pack_forget()
            new_risk_area.container_frame.pack_forget()
            pack_main_area()

        elif not new_risk_area.container_frame.winfo_ismapped():
            # Displays the new_risk_area if hidden
            print("new_risk_area is not visible, packing new_risk_area")

            new_risk_area.container_frame.pack(fill=BOTH, expand=True)

    def current_assessments_area_control():
        """Opens and closes the current_assesments_area based on
        whether or not it is already packed
        """
        print("Triggered current_assessments_area_control()")
        new_risk_area.container_frame.pack_forget()
        risk_assessment_area.container_frame.pack_forget()
        main_title.pack_forget()
        intro_text.pack_forget()
        current_assessments_area.container_frame_scroll.pack_forget()

        # Hides the new_risk_area if packed
        if current_assessments_area.container_frame.winfo_ismapped():
            print("new_risk_area is visible, hiding current_assesments_area")
            current_assessments_area.container_frame.pack_forget()
            pack_main_area()

        # Displays the new_risk_area if hidden
        elif not current_assessments_area.container_frame.winfo_ismapped():
            print("new_risk_area is not visible, packing side_bar")
            current_assessments_area.container_frame.pack(
                fill=BOTH, expand=True, side=TOP
            )

    def risk_assessment_area_control():
        """Opens and closes the risk_assesment_area based on
        whether or not it is already packed
        """
        print("Triggered risk_assessments_area_control()")

        # Hides the new_risk_area if packed
        if risk_assessment_area.container_frame.winfo_ismapped():
            print("risk_assessment_area is visible, hiding risk_assesment_area")
            risk_assessment_area.container_frame.pack_forget()
            main_title.pack(expand=True, side=TOP, anchor=N, fill=X)
            intro_text.pack(expand=True, side=TOP, anchor=N, fill=X)

        # Displays the new_risk_area if hidden
        elif not risk_assessment_area.container_area.winfo_ismapped():
            print("risk_assessment_area is not visible, packing risk_assessment_area")
            for widget in main_area.winfo_children():
                widget.pack_forget()

            risk_assessment_area.container_frame.pack(fill=BOTH, expand=True, side=TOP)
            risk_assessment_area.container_area.pack_configure(expand=True, anchor=N)

    def pack_risk_area():
        risk_assessment_title.pack(
            fill=X, expand=False, side=TOP, anchor=N, padx=15, pady=5
        )
        risk_assessment_name.pack(expand=False, side=TOP, anchor=N, padx=15, pady=5)
        risk_assessment_datetime_time.pack(
            fill=X, expand=False, side=TOP, anchor=N, padx=25, pady=5
        )

        risk_assessment_area_explanation.pack(side=TOP, anchor=N, fill=X)
        risk_assessment_area_explanation.insert(
            "1.5",
            """
            Risk Level is based on likelihood of incidence:  
            0 = impossible 1 = unlikely, 2 = possible, 3 = likely, 4 probable, 5 = certain, or N/A = Not Applicable. 
            
            This is multiplied by the severity of injuries likely : 
            1 = trivial, 2 = minor, 3 = severe, 4 = major, 5 = fatal. The resulting number quantifies the overall risk.  
            
            Higher number = higher priority. 

            To change the values of a risk, select the row in the table below and update the values on the right side of the
            screen before pressing Enter.
            """,
        )
        risk_assessment_area_explanation.configure(
            state=DISABLED,
            bg=med_grey,
            fg=white,
            borderwidth=0,
            font="lucida 12",
        )

        container_box.pack(side=RIGHT, fill=Y)

        risk_assessment_info_container_table.pack(
            fill=BOTH, expand=3, anchor=N, padx=5, pady=5, side=LEFT
        )

        risk_table_scroll.pack(side=RIGHT, fill=Y)

        """update_title.pack(side=TOP, expand=False, pady=10, anchor=W)"""
        update_risk_label.pack(side=TOP, expand=False, pady=3, anchor=W)
        update_risk.pack(side=TOP, expand=False, pady=3, anchor=W)
        update_likelihood_label.pack(side=TOP, expand=False, pady=3, anchor=W)
        update_likelihood.pack(side=TOP, expand=False, pady=3, anchor=W)
        update_severity_label.pack(side=TOP, expand=False, pady=3, anchor=W)
        update_severity.pack(side=TOP, expand=False, pady=3, anchor=W)
        update_actions_label.pack(side=TOP, expand=False, pady=3, anchor=W)
        update_actions.pack(side=TOP, expand=False, pady=3, anchor=W)
        update_actions_submit.pack(side=TOP, expand=False, pady=3, anchor=N)

    def delete_risks(x):
        """Destroys risk modules widget that is right clicked
        in the side bar.
        Also deletes the related file from file_location_risks_folder.
        Finally, refreshes the side bar to reflect the current
        risk modules.
        """
        print("Triggered delte_risk(x)")
        delete_risk_query = messagebox.askyesno(
            "Delete risk", "Delete this risk module?"
        )
        if delete_risk_query == True:
            risk_modules = [x for x in os.listdir(file_location_risks_folder)]
            if f"{x}" in risk_modules:
                print(f"Deleting {x}")
                os.chdir(file_location_risks_folder)
                os.remove(x)
                os.chdir(file_location)

        else:
            print(f"Cancelled deleting {x}")

        refresh_sidebar()

    def home():
        """Unpacks all current widgets from the main_area.
        Packs the main area with title and information to
        serve as a home screen.
        """
        print("Triggered home()")
        print("Hiding all packed widgets")
        table_contents.clear()
        for widget in main_area.winfo_children():
            widget.pack_forget()
        for widget in risk_assessment_area.container_area.winfo_children():
            widget.pack_forget()
        pack_main_area()
        pack_risk_area()

    def save():
        """Retreieves information displayed in the risk_assesment_area
        and generates a new .docx file in the risk_assessments_folder
        """
        print("Triggered save()")
        save_query = messagebox.askyesno(
            "Save", "Would you like to save this risk assessment?"
        )
        if save_query == True:
            print("Saving file")
            file_title = risk_assessment_name.get().rstrip().title()
            cur_date = f"{now.day}.{now.month}.{now.year}"
            file_title_path = os.path.join(
                risk_assessments_folder, file_title + f" {cur_date}.txt"
            )

            if file_title == "Enter Title" or None:
                pass

            else:
                os.chdir(templates_folder)
                document = Document()
                records = []

                risk_details = ""
                with open(
                    os.path.join(templates_folder, "risk_assessment_opening_para.txt"),
                    "r",
                ) as template:
                    for i in risk_assessment_info_container_table.get_children():
                        records.append(
                            risk_assessment_info_container_table.item(i)["values"]
                        )

                    starting_title = document.add_heading(
                        "St. John the Baptist\nRisk Assessment Form", 0
                    )

                    title_heading = document.add_heading(
                        f"Risk Assessment: {risk_assessment_name.get().title()}", 1
                    )
                    sub_heading = document.add_heading(
                        f"Date of assessment: {now.date()}", 1
                    )

                    paragraph = document.add_paragraph(template.read())

                    document.add_page_break()

                    risk_details_para = document.add_paragraph(
                        """Risk Level is based on likelihood of incidence:  
                    0 = impossible 1 = unlikely, 2 = possible, 3 = likely, 4 probable, 5 = certain, or N/A = Not Applicable. 

                    This is multiplied by the severity of injuries likely : 
                    1 = trivial, 2 = minor, 3 = severe, 4 = major, 5 = fatal. The resulting number quantifies the overall risk.  
                    Higher number = higher priority. 
                        """
                    )

                    risk_table_doc = document.add_table(rows=1, cols=5)
                    risk_table_doc.style = "Table Grid"

                    hdr_cells = risk_table_doc.rows[0].cells
                    hdr_cells[0].text = "Risk/Hazard"
                    hdr_cells[1].text = "Likelihood"
                    hdr_cells[2].text = "Severity"
                    hdr_cells[3].text = "Outcome"
                    hdr_cells[4].text = "Actions"

                    for risk, likelihood, severity, outcome, actions in records:
                        row_cells = risk_table_doc.add_row().cells
                        row_cells[0].text = str(risk)
                        row_cells[1].text = str(likelihood)
                        row_cells[2].text = str(severity)
                        row_cells[3].text = str(outcome)
                        row_cells[4].text = str(actions)
                    document_file_path = os.path.join(
                        risk_assessments_folder,
                        f"{risk_assessment_name.get().title().rstrip()} - {cur_date}.docx",
                    )

                    document.save(document_file_path)

                    for child in risk_assessment_info_container_table.get_children():
                        risk_assessment_info_container_table.delete(child)
                        table_contents.clear()
                        risk_assessment_name.config(state=NORMAL)
                    risk_assessment_name.delete(0, END)
                    risk_assessment_name.insert(0, "Enter Title")

                    template.seek(0)
                    template.close()

                def display_risk_files(*args):
                    """Dynamically creates and wraps buttons
                    for opening .docx files created by the user.
                    Buttons are packed into the current_assessments_area.
                    """
                    global wrap_widget
                    wrap_widget += 1
                    risk_assessment_files = Button(
                        current_assessments_area_wrap_buttons,
                        text=new_file.replace(".docx", ""),
                        font="lucida 10 bold",
                        width=37,
                        height=1,
                        command=lambda new_file=new_file: open_risk_assessment(
                            new_file
                        ),
                        bg=med_grey,
                        fg=white,
                        borderwidth=1,
                        highlightthickness=0,
                        activebackground=med_grey,
                        activeforeground=white,
                        pady=10,
                    )
                    risk_assessment_files.pack(side=LEFT)

                    current_assessments_area_wrap_buttons.configure(state="normal")
                    current_assessments_area_wrap_buttons.window_create(
                        "insert", window=risk_assessment_files, padx=10, pady=10
                    )
                    current_assessments_area_wrap_buttons.configure(state=DISABLED)

                for widget in current_assessments_area_wrap_buttons.winfo_children():
                    widget.destroy()
                for new_file in os.listdir(
                    os.path.join(file_location, "Risk Assessments")
                ):
                    display_risk_files(new_file)

                os.chdir(file_location)
                messagebox.showinfo(
                    "Saved file",
                    f"Saved file as '{risk_assessment_name.get().title()}'",
                )
                refresh_sidebar()
        else:
            print("Cancelled save")
            pass

    def pack_main_area():
        """Packs the main title and introductory
        text into the main area
        """
        main_title.pack(expand=False, side=TOP, anchor=N, fill=X)
        intro_text.pack(expand=False, side=TOP, anchor=N, fill=X)

    """# search function --------------------------------UNASSIGNED----------
    def search():
        print("Triggered search()")
        pass"""

    # about function ---------------------------------UNASSIGNED----------
    def about():
        """Displays information about the program"""
        print("Triggered about()")
        pass

    # kill function
    def kill():
        """Displays a message box that querys the user
        before closing the program.
        """
        if messagebox.askokcancel("Quit", "Do you want to quit?"):
            master_window.destroy()

    def update_table(*args):
        """Updates the risk_assessment_info_container_table row
        values based on user input.
        """
        print("Triggered update_table()")

        def numbers_only():
            """ """
            messagebox.showinfo(
                "User error",
                "Likelihood must be a number.\n\n Number must be between 1-5.",
            )
            pass

        try:
            int(update_likelihood.get())

        except ValueError:
            print(f"Cannot convert {update_likelihood.get()} to integer.")
            numbers_only()
            return

        try:
            int(update_severity.get())
        except:
            print(f"Cannot convert {update_severity.get()} to integer.")
            numbers_only()
            return

        if not int(update_likelihood.get()) in range(0, 5) or not int(
            update_severity.get()
        ) in range(0, 5):
            numbers_only()

        else:
            row_id = risk_assessment_info_container_table.selection()
            table_data = risk_assessment_info_container_table.item(row_id)["values"]

            selected = risk_assessment_info_container_table.focus()

            updated_outcome = int(update_likelihood.get().rstrip()) * int(
                update_severity.get().rstrip()
            )

            def wrap(string, length=60):
                """Wraps widgets that go beyond the width of the screen."""
                if len(update_actions.get("1.0", END)) >= 30:
                    return "\n".join(textwrap.wrap(string, length))

                else:
                    return update_actions.get("1.0", END)

            risk_range = range(0 - 6)

            risk_assessment_info_container_table.item(
                selected,
                text="",
                values=(
                    update_risk.get().rstrip().title(),
                    update_likelihood.get().rstrip(),
                    update_severity.get().rstrip(),
                    updated_outcome,
                    wrap(update_actions.get("1.0", END)),
                ),
            )

            update_risk.delete(0, END)
            update_likelihood.delete(0, END)
            update_severity.delete(0, END)
            update_actions.delete("1.0", END)
            update_risk.focus_set()

    def display_risk(fn_name):
        """Adds values to the rows in
        risk_assessment_info_container_table that
        are later retrieved for saving into a .docx file.
        """

        print("Triggered display_risk()")
        print(table_contents)
        main_title.pack_forget()
        intro_text.pack_forget()
        current_assessments_area.container_frame.pack_forget()
        new_risk_area.container_frame.pack_forget()
        risk_assessment_area.container_frame.pack_configure(expand=True)
        risk_assessment_area.container_area.pack_configure(expand=True, anchor=N)

        risk_file = os.path.join(file_location_risks_folder, fn_name)

        risk_assessment_area.container_canvas.configure(
            scrollregion=risk_assessment_area.container_canvas.bbox("all")
        )

        risk_assessment_area.container_frame_scroll.pack(
            fill=Y, expand=False, side=RIGHT
        )
        if risk_assessment_area.container_frame.winfo_ismapped():
            pass
        elif not risk_assessment_area.container_frame.winfo_ismapped():
            risk_assessment_area.container_frame.pack(fill=BOTH, expand=True, side=TOP)

        if fn_name in table_contents:
            print(f"{fn_name} widget found.")
            for id in risk_assessment_info_container_table.get_children():
                print(risk_assessment_info_container_table.item(id)["values"])
                if fn_name in risk_assessment_info_container_table.item(id)["values"]:
                    print(f"Found {id}")
                    risk_assessment_info_container_table.delete([int(id)])
                else:
                    ("Not found")
            table_contents.remove(fn_name)
            print(table_contents)

        else:
            new_risk_name = new_risk_area_entry_one.get().rstrip().title()
            new_risk_area_entry_one.delete(0, "end")
            new_risk_likeliehood = 1
            new_risk_severity = 1
            new_risk_outcome = int(new_risk_likeliehood) * int(new_risk_severity)
            new_risk_actions = "N/A"

            print(f"{fn_name} widget not found, packing {fn_name}")
            # Checks for existing risk_file, if not found will create a new json file.
            if not os.path.exists(risk_file):
                print(f"File not found. Creating new file {risk_file} display_risk()")
                with open(risk_file, "w+") as risk_file_json:
                    write_data = {
                        f"Risk": f"{new_risk_name}",
                        "Likelihood": f"{new_risk_likeliehood}",
                        "Severity": f"{new_risk_severity}",
                        "Outcome": f"{new_risk_outcome}",
                        "Actions": f"{new_risk_actions}",
                    }
                    risk_file_dump = json.dumps(write_data, indent=4)
                    risk_file_json.write(risk_file_dump)
                    risk_file_json.seek(0)
                    read_data = risk_file_json.read()
                    risk_file_json.close()

            else:
                with open(risk_file, "r+") as risk_file_json:

                    print(f"{risk_file} found")
                    read_data = risk_file_json.read()
                    risk_file_json.close()
                    data_text = json.loads(read_data)

                    risk_text = data_text["Risk"]
                    likelihood_text = data_text["Likelihood"]
                    severity_text = data_text["Severity"]
                    outcome_text = int(likelihood_text) * int(severity_text)
                    actions_text = data_text["Actions"]

                    next_row = next(table_row)

                    risk_assessment_info_container_table.insert(
                        parent="",
                        index="end",
                        iid=next_row,
                        text="",
                        values=(
                            risk_text,
                            likelihood_text,
                            severity_text,
                            outcome_text,
                            actions_text,
                        ),
                    )

                    table_contents.append(fn_name)

    # endregion

    # region Classes

    # --------------------------------CLASSES---------------------------------
    # Risks Class
    class Risks:
        """Class for dynamically creating
        buttons that serve as modules for
        risks created by the user.
        Risks are packed into the side bar.
        """

        def __init__(self, window, text, row, column, command=None) -> None:

            self.window = window
            self.text = text
            self.row = row
            self.column = column
            self.command = command

            new_risks_button = Button(
                window,
                text=self.text,
                width=15,
                height=4,
                command=self.command,
                bg=light_grey,
                fg=white,
                borderwidth=0,
                highlightthickness=0,
                activebackground=med_grey,
                activeforeground=white,
                font="lucidia 9",
            )
            new_risks_button.pack(padx=1, pady=1)
            new_risks_button.bind(
                "<Button-3>", lambda x=self.text: delete_risks(self.text)
            )

    class Icons:
        """Class for creating buttons to serve
        as icons that are packed into the icon bar.
        """

        def __init__(self, window, text, row, column, command=None) -> None:

            self.window = window
            self.text = text
            self.row = row
            self.column = column
            self.command = command

            new_icon = Button(
                window,
                text=self.text,
                font="lucida 12 bold",
                width=6,
                height=2,
                command=self.command,
                bg=light_grey,
                fg=white,
                borderwidth=0,
                highlightthickness=0,
                activebackground=med_grey,
                activeforeground=white,
            )
            new_icon.grid(row=self.row, column=self.column, padx=1, pady=0)

    class FileBarButtons:
        """Class used to create buttons to serve as icons
        that are packed into the app_menu
        """

        def __init__(self, window, text, command=None) -> None:

            self.window = window
            self.text = text
            self.command = command

            # Creates a new button intended to
            # serve as an icon on the icon_bar.
            new_icon = Button(
                window,
                text=self.text,
                font="lucida 10 bold",
                command=self.command,
                width=10,
                height=2,
                bg=light_grey,
                fg=white,
                borderwidth=0,
                highlightthickness=0,
                activebackground=med_grey,
                activeforeground=white,
            )
            new_icon.pack(side=RIGHT)

    class ScrollableFrame:
        """Class for creating frames that appear to scroll
        the widgets that are inside of them.
        """

        def __init__(self, master_window):

            self.master_window = master_window

            self.container_frame = LabelFrame(
                master_window,
                width=200,
                padx=1,
                pady=1,
                relief=SUNKEN,
                bg=light_grey,
                fg=white,
                borderwidth=0,
                highlightthickness=0,
            )
            self.container_frame.pack(fill=BOTH, expand=False, side=LEFT)

            self.container_canvas = Canvas(
                self.container_frame,
                width=100,
                bg=med_grey,
                borderwidth=0,
                highlightthickness=0,
            )
            self.container_canvas.pack(fill=BOTH, expand=1, side=LEFT)

            self.container_frame_scroll = Scrollbar(
                self.container_frame,
                bg=light_grey,
                orient=VERTICAL,
                troughcolor=dark_grey,
                background=dark_grey,
                command=self.container_canvas.yview,
            )
            self.container_frame_scroll.pack(side=RIGHT, fill=Y, padx=1)

            self.container_area = LabelFrame(
                self.container_canvas,
                width=200,
                padx=1,
                pady=1,
                relief=SUNKEN,
                bg=med_grey,
                fg=white,
                borderwidth=0,
                highlightthickness=0,
            )
            self.container_area.pack(fill=BOTH, expand=False, side=LEFT)

            self.container_canvas.create_window(
                (0, 0), window=self.container_area, anchor=N + W
            )

            self.container_canvas.configure(
                yscrollcommand=self.container_frame_scroll.set
            )
            self.container_canvas.bind(
                "<Configure>",
                lambda e: self.container_canvas.configure(
                    scrollregion=self.container_canvas.bbox("all")
                ),
            )

    # endregion

    # region Main Script
    """Top Menu Bar Label Frame (app_menu)
    A replacement menu (frame) that sits
    at the top of the program (app_menu)"""

    # region App Menu
    app_menu = LabelFrame(
        master_window,
        width=50,
        height=35,
        padx=1,
        pady=1,
        relief=None,
        bg=light_grey,
        fg=white,
        borderwidth=0,
        highlightthickness=0,
    )
    app_menu.pack(fill=X, expand=False, side=TOP)

    # Creates a row of buttons on the (app_menu)
    row = 0
    commands = [kill, about, save]  # remove start search
    command_names = ["Exit", "About", "Save"]  # "Remove" "Start" "Search"
    num_of_buttons = len(commands)

    for i in range(num_of_buttons):
        my_command = commands.pop(0)
        my_command_names = command_names.pop(0)
        while row <= num_of_buttons:
            row += 1
            new_row = [].append(row)
            new_button = FileBarButtons(app_menu, my_command_names, my_command)
            break
    # endregion

    # region Icon Bar
    """Icon Bar Label Frame (icon_bar)
    A frame that packs on the left of the
    screen that contains icons and is always visible"""
    icon_bar = LabelFrame(
        master_window,
        width=50,
        padx=1,
        pady=1,
        relief=None,
        bg=light_grey,
        fg=white,
        borderwidth=0,
        highlightthickness=0,
    )
    icon_bar.pack(fill=BOTH, expand=False, side=LEFT)

    """Functions that will be pushed into the dynamically
    created icon bar.
    """
    icon_bar_fns = [
        home,
        new_risk_area_control,
        risk_assessment_area_control,
        current_assessments_area_control,
        clear_risks,
        side_bar_control,
    ]

    # Variables for determening how many icons to generate
    icon_bar_icons_name = ["H", "+", "R", "C", "X", "S"]
    icon_bar_icons_name_const = ["H", "+", "R", "C", "X", "S"]
    icon_row = 0
    num_of_icon_bar_icons = len(icon_bar_fns)

    for i in range(num_of_icon_bar_icons):
        """Creates icons in the icon bar"""
        icon_bar_icons_pop = icon_bar_fns.pop(0)
        icon_bar_icons_name_pop = icon_bar_icons_name.pop(0)
        while icon_row <= len(icon_bar_icons_name_const):
            icon_row += 1
            new_row = [].append(icon_row)
            new_button = Icons(
                icon_bar, icon_bar_icons_name_pop, icon_row, 0, icon_bar_icons_pop
            )
            break

    # endregion

    # region Side bar

    side_bar = ScrollableFrame(master_window)
    """Side Bar Label Frame (side_bar)
    A frame that packs on the left of the screen
    that is toggleable and contains buttons for
    inputting pre-written modules"""
    # side_bar.container_frame.pack_forget()

    side_bar.container_frame.pack_forget()

    refresh_sidebar()

    def new_risk(*args):
        """Allows the user to create new risk modules
        that will be stored in the "Risks" directory
        and packed into the side_bar.
        Side bar is then refreshed to reflect the
        newly created risks
        """
        print("Triggered new_risk()")
        new_risk_area_entry_one.focus_set()

        new_risk_name = new_risk_area_entry_one.get().rstrip().title()
        # Creates default values for the new risk
        new_risk_area_entry_one.delete(0, "end")
        new_risk_likeliehood = 1
        new_risk_severity = 1
        new_risk_outcome = int(new_risk_likeliehood) * int(new_risk_severity)
        new_risk_actions = "N/A"

        if not new_risk_name == "":
            new_risk_file = os.path.join(file_location_risks_folder, new_risk_name)
            # Checks for existing risk_file, if not found will create a new json file
            if not os.path.exists(new_risk_file):
                print(f"File not found. Creating new file {new_risk_name}")

                """Creates a new file and writes default data
                before saving to risks directory."""
                with open(new_risk_file, "w+") as new_risk_file_json:
                    write_data = {
                        f"Risk": f"{new_risk_name}",
                        "Likelihood": f"{new_risk_likeliehood}",
                        "Severity": f"{new_risk_severity}",
                        "Outcome": f"{new_risk_outcome}",
                        "Actions": f"{new_risk_actions}",
                    }
                    new_risk_file_dump = json.dumps(write_data, indent=4)
                    new_risk_file_json.write(new_risk_file_dump)
                    new_risk_file_json.close()

                    refresh_sidebar()

    # endregion

    # region Main area

    """Main Area Label Frame (main_area)
    Creates a frame that contains other
    frames that can be toggled open or closed.
    Serves as a home screen when other
    frames have been unpacked/hidden.
    """
    main_area = LabelFrame(
        master_window,
        width=200,
        padx=1,
        pady=1,
        relief=SUNKEN,
        bg=dark_grey,
        fg=white,
        borderwidth=0,
        highlightthickness=0,
    )
    main_area.pack(fill=BOTH, expand=True, side=RIGHT)

    # Title displayed in the main_area homescreen
    main_title = Label(
        main_area,
        text="""
Risk Assesment
Builder
        """,
        font="lucida 40 bold",
        bg=dark_grey,
        pady=15,
        padx=25,
        fg=white,
        borderwidth=0,
    )

    # Usage information displayed in the main_area homescreen
    intro_text = Label(
        main_area,
        text="""
Welcome to the Risk Assessment Builder.

To being adding a new Risk Module to the modules side bar
press '+' on the icon bar.

To get started building your risk assessment
select 'R' on the icon bar. 

To search for and open saved risk assessments press 'C'.

Press '+' on the icon bar to add new risk modules to the modules side bar.

To clear all current risks press 'X' on the icon bar.

To hide the side bar press 'S' on the icon bar.
        """,
        font="lucida 15",
        bg=dark_grey,
        pady=20,
        padx=25,
        fg=white,
        borderwidth=0,
        justify=CENTER,
    )
    pack_main_area()

    # endregion

    # region New Risk Area
    new_risk_area = ScrollableFrame(main_area)
    """New Risk Area
    A scrollable frame that packs on the right of the screen
    that is toggleable and contains widgets for
    inputting new risk modules"""
    new_risk_area.container_frame_scroll.pack_forget()
    new_risk_area.container_area.pack_configure(fill=BOTH, expand=True)
    new_risk_area.container_frame.pack_forget()

    # Title at the top of Current Assessments Label Frame
    new_risk_area_title_label = Label(
        new_risk_area.container_area,
        font="lucida 40 bold",
        text="""
Add New Risk
        """,
        bg=med_grey,
        fg=white,
        justify=LEFT,
    )
    new_risk_area_title_label.pack(
        fill=None, expand=False, side=TOP, anchor=N, padx=15, pady=20
    )

    # New Risk Area Subframe (new_risk_area)
    new_risk_area_subframe = LabelFrame(
        new_risk_area.container_area,
        width=200,
        padx=1,
        pady=1,
        relief=SUNKEN,
        bg=med_grey,
        fg=white,
        borderwidth=0,
        highlightthickness=0,
    )
    new_risk_area_subframe.pack(fill=X, expand=False)

    new_risk_area_entry_one_label = Label(
        new_risk_area_subframe,
        font="lucida 20 bold",
        text="Enter hazard:\ne.g. 'Fire'",
        bg=med_grey,
        fg=white,
    )
    new_risk_area_entry_one_label.pack(
        fill=None, expand=False, padx=15, pady=5, anchor=N
    )

    new_risk_area_entry_one = Entry(
        new_risk_area_subframe,
        textvariable=1,
        font="lucida 10 bold",
        width=21,
        bg=white,
        fg=med_grey,
        borderwidth=1,
        highlightthickness=0,
    )
    new_risk_area_entry_one.pack(expand=False, pady=10, anchor=N)
    new_risk_area_entry_one.bind("<Return>", new_risk)

    risk_submission = Button(
        new_risk_area_subframe,
        text="Submit",
        font="lucida 12 bold",
        width=15,
        height=1,
        command=new_risk,
        bg=dark_grey,
        fg=white,
        borderwidth=1,
        highlightthickness=0,
        activebackground=med_grey,
        activeforeground=white,
        pady=10,
    )
    risk_submission.pack(expand=False, pady=10, anchor=N)

    new_risk_area_explanation = Label(
        new_risk_area.container_area,
        font="lucida 15",
        text="""
This area is used to add new risks to the
side bar that are likely to occur regularly.

Once added, the side bar can be used to add
re-occuring risks to a new risk assesment template.
    """,
        bg=med_grey,
        fg=white,
        justify=CENTER,
        pady=20,
    )

    new_risk_area_explanation.pack(side=TOP, anchor=N, fill=X)
    # endregion

    # region Current Assesments Area
    current_assessments_area = ScrollableFrame(main_area)
    """Current Assessments Area (current_assessments_area)
    creates a toggleable frame that is scrollable and contains
    risk assessments that have been created and stored in "Risks"
    Directory.
    """

    current_assessments_area.container_frame.pack_forget()
    current_assessments_area.container_frame_scroll.pack_forget()
    current_assessments_area.container_area.pack(
        side=TOP, expand=1, fill=BOTH, pady=10, anchor=W
    )

    current_assessment_area_main_title = Label(
        current_assessments_area.container_area,
        text="""
    Current Risk Assessments
        """,
        font="lucida 40 bold",
        bg=med_grey,
        pady=15,
        padx=25,
        fg=white,
        borderwidth=0,
    )
    current_assessment_area_main_title.pack(expand=True, side=TOP, anchor=N, fill=X)

    current_assessments_area_wrap_buttons = ScrolledText(
        current_assessments_area.container_area,
        font="lucida 12 bold",
        bg=med_grey,
        fg=med_grey,
        borderwidth=0,
        highlightthickness=0,
        wrap=WORD,
        relief=SUNKEN,
        yscrollcommand=lambda *args: current_assessments_area.container_frame_scroll.set(
            *args
        ),
    )
    current_assessments_area_wrap_buttons.pack(
        side=TOP, expand=1, fill=BOTH, pady=10, anchor=W
    )

    current_assessments_area.container_frame_scroll.config(
        command=current_assessments_area_wrap_buttons.yview
    )

    if os.listdir(os.path.join(file_location, "Risk Assessments")) == 0 or None:
        print("No files to display.")
        pass
    else:

        def delete_file(event):
            """Destroys the widget that triggered the function and
            removes the associated file from the "Risk Assessments"
            directory.
            """
            print("Triggered delete_file(event)")
            delete_file_query = messagebox.askyesno(
                "Delete risk assessment", "Delete this risk assessment?"
            )

            if delete_file_query == True:
                # print(event)
                btn_data = f'{event.widget.cget("text")}.docx'
                print(event.widget.cget("text"))
                os.chdir(risk_assessments_folder)
                os.remove(str(btn_data))
                os.chdir(file_location)
                event.widget.destroy()
                print("Widget destroyed")

            else:
                pass

        def add_widget(*args):
            """Packs buttons onto the current_assessment_area
            that wrap to fit inside the screen.
            Buttons run functions to open .docx files with
            the same name that are stored in the "Risk Assessments"
            directory that have been created by the user.
            """
            global wrap_widget
            wrap_widget += 1
            risk_assessment_files = Button(
                current_assessments_area_wrap_buttons,
                text=new_file.replace(".docx", ""),
                font="lucida 10 bold",
                width=38,
                height=1,
                command=lambda new_file=new_file: open_risk_assessment(new_file),
                bg=med_grey,
                fg=white,
                borderwidth=1,
                highlightthickness=0,
                activebackground=med_grey,
                activeforeground=white,
                pady=10,
            )
            risk_assessment_files.pack(side=LEFT)
            risk_assessment_files.bind(
                "<Button-3>", lambda new_file=new_file: delete_file(new_file)
            )

            current_assessments_area_wrap_buttons.configure(state=NORMAL)
            current_assessments_area_wrap_buttons.window_create(
                "insert", window=risk_assessment_files, padx=10, pady=10
            )
            current_assessments_area_wrap_buttons.configure(state=DISABLED)

        for new_file in os.listdir(os.path.join(file_location, "Risk Assessments")):
            add_widget()

    # endregion

    # region Risk Assessment Area

    risk_assessment_area = ScrollableFrame(main_area)
    """Risk Assessment Area (risk_assessment_area)
    creates a toggleable frame that is scrollable and contains
    a table and widgets for building risk assessments that
    are later retrieved and saved.
    """
    risk_assessment_area.container_frame_scroll.pack_forget()
    risk_assessment_area.container_area.pack_configure(fill=BOTH)
    risk_assessment_area.container_frame.pack_forget()

    risk_assessment_title = Label(
        risk_assessment_area.container_area,
        font="lucida 40 bold",
        text="Risk Assessment",
        bg=med_grey,
        fg=white,
    )
    risk_assessment_name = Entry(
        risk_assessment_area.container_area,
        font="lucida 16 bold",
        bg=med_grey,
        fg=white,
        justify=CENTER,
        borderwidth=0,
        insertbackground=white,
        width=40,
    )
    risk_assessment_datetime_time = Label(
        risk_assessment_area.container_area,
        font="lucida 16 bold",
        text=f"{now.ctime()}",
        bg=med_grey,
        fg=white,
        justify=CENTER,
    )
    risk_assessment_area_explanation = Text(
        risk_assessment_area.container_area,
        font="lucida 12 bold",
        height=12,
        bg=med_grey,
        fg=white,
        borderwidth=0,
        highlightthickness=0,
        wrap=WORD,
        relief=SUNKEN,
        highlightbackground=med_grey,
    )

    container_box = LabelFrame(
        risk_assessment_area.container_area,
        width=400,
        padx=20,
        pady=1,
        relief=SUNKEN,
        bg=med_grey,
        fg=white,
        borderwidth=0,
        highlightthickness=0,
    )

    risk_assessment_name.insert(0, "Enter Title")
    risk_assessment_name.bind(
        "<Return>",
        lambda i=i: risk_assessment_name.config(
            state=DISABLED,
            disabledbackground=med_grey,
            disabledforeground=white,
            fg=white,
            borderwidth=0,
            font="lucida 14 bold",
        ),
    )
    risk_assessment_name.bind(
        "<FocusIn>", lambda i=i: risk_assessment_name.delete(0, END)
    )

    def focus_out(e):
        if risk_assessment_name.get() == "" or None:
            risk_assessment_name.insert(0, "Enter Title")
        else:
            pass

    risk_assessment_name.bind("<FocusOut>", focus_out)

    risk_table_scroll = Scrollbar(risk_assessment_area.container_area)

    risk_assessment_info_container_table = ttk.Treeview(
        risk_assessment_area.container_area, yscrollcommand=risk_table_scroll.set
    )
    risk_assessment_info_container_table["columns"] = (
        "Risk",
        "Likelihood",
        "Severity",
        "Outcome",
        "Actions",
    )

    risk_assessment_info_container_table.column("#0", width=0, stretch=NO)
    risk_assessment_info_container_table.column("Risk", anchor=CENTER, width=80)
    risk_assessment_info_container_table.column("Likelihood", anchor=CENTER, width=20)
    risk_assessment_info_container_table.column("Severity", anchor=CENTER, width=20)
    risk_assessment_info_container_table.column("Outcome", anchor=CENTER, width=20)
    risk_assessment_info_container_table.column("Actions", anchor=CENTER, width=400)

    risk_assessment_info_container_table.heading("#0", text="", anchor=CENTER)
    risk_assessment_info_container_table.heading("Risk", text="Risk", anchor=CENTER)
    risk_assessment_info_container_table.heading(
        "Likelihood", text="Likelihood", anchor=CENTER
    )
    risk_assessment_info_container_table.heading(
        "Severity", text="Severity", anchor=CENTER
    )
    risk_assessment_info_container_table.heading(
        "Outcome", text="Outcome", anchor=CENTER
    )
    risk_assessment_info_container_table.heading(
        "Actions", text="Actions", anchor=CENTER
    )

    risk_table_scroll.config(command=risk_assessment_info_container_table.yview)
    risk_assessment_info_container_table.bind("<ButtonRelease-1>", edit_table)

    update_title = Label(
        container_box,
        font="lucida 14 bold",
        text="Update Risk",
        pady=1,
        padx=5,
        bg=med_grey,
        fg=white,
    )

    update_risk = Entry(
        container_box,
        textvariable=4,
        font="lucida 12 bold",
        width=21,
        bg=white,
        fg=med_grey,
        borderwidth=1,
        highlightthickness=0,
        relief=SUNKEN,
    )
    update_risk_label = Label(
        container_box,
        font="lucida 12 bold",
        text="Risk",
        pady=1,
        padx=5,
        bg=med_grey,
        fg=white,
    )
    update_likelihood = Entry(
        container_box,
        textvariable=5,
        font="lucida 12 bold",
        width=21,
        bg=white,
        fg=med_grey,
        borderwidth=1,
        highlightthickness=0,
        relief=SUNKEN,
    )
    update_likelihood_label = Label(
        container_box,
        font="lucida 12 bold",
        text="Likelihood",
        pady=1,
        padx=5,
        bg=med_grey,
        fg=white,
    )
    update_severity = Entry(
        container_box,
        textvariable=6,
        font="lucida 12 bold",
        width=21,
        bg=white,
        fg=med_grey,
        borderwidth=1,
        highlightthickness=0,
        relief=SUNKEN,
    )
    update_severity_label = Label(
        container_box,
        font="lucida 12 bold",
        text="Severity",
        pady=1,
        padx=5,
        bg=med_grey,
        fg=white,
    )
    update_actions = Text(
        container_box,
        font="lucida 12 bold",
        width=40,
        bg=white,
        fg=med_grey,
        borderwidth=1,
        highlightthickness=0,
        wrap=WORD,
        relief=SUNKEN,
    )
    update_actions_label = Label(
        container_box,
        font="lucida 12 bold",
        text="Actions",
        pady=1,
        padx=5,
        bg=med_grey,
        fg=white,
    )

    update_actions_submit = Button(
        container_box,
        text="Submit",
        font="lucida 12 bold",
        width=19,
        height=1,
        command=update_table,
        bg=dark_grey,
        fg=white,
        borderwidth=1,
        highlightthickness=0,
        activebackground=med_grey,
        activeforeground=white,
        pady=10,
    )
    update_risk.bind("<Return>", update_table)
    update_likelihood.bind("<Return>", update_table)
    update_severity.bind("<Return>", update_table)

    # endregion

    pack_risk_area()

    # Defines the master_window attributes
    master_window.title("Risk Assessment Builder")
    master_window.geometry("800x600")

    # Initializes the program
    master_window.mainloop()
    # endregion


if __name__ == "__main__":
    main()
